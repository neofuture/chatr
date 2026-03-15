#!/usr/bin/env python3
"""
Generate the Chatr professional pitch + technical reference.
Run with: python3 scripts/generate-presentation-docx.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), '..', 'Chatr - Product Overview.docx')
SS  = os.path.join(os.path.dirname(__file__), '..', 'screenshots')

NAVY    = RGBColor(0x0F, 0x17, 0x2A)
BLUE    = RGBColor(0x3B, 0x82, 0xF6)
SLATE   = RGBColor(0x33, 0x41, 0x55)
GREY    = RGBColor(0x64, 0x74, 0x8B)
ACCENT  = RGBColor(0xF9, 0x73, 0x16)

MOBILE_W = Inches(2.5)       # mobile screenshots (portrait ~1:2 ratio → ~5" tall = half page)
WIDE_W   = Inches(5.0)       # dashboard / wide screenshots
PAIR_W   = Inches(2.3)       # each image in a side-by-side pair

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.2)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = SLATE
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# ── Helpers ──────────────────────────────────────────────────────────────────
_sn = [0]; _sub = [0]

def _keep(p):
    """Prevent paragraph from being separated from the next one across pages."""
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))

def H(text, level=1, numbered=True):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        if numbered: _sn[0] += 1; _sub[0] = 0
        prefix = f'{_sn[0]}. ' if numbered else ''
        r = h.add_run(prefix + text)
        r.font.size = Pt(18); r.font.color.rgb = NAVY; r.bold = True
        h.paragraph_format.space_before = Pt(14); h.paragraph_format.space_after = Pt(4)
    elif level == 2:
        if numbered: _sub[0] += 1; prefix = f'{_sn[0]}.{_sub[0]} '
        else: prefix = ''
        r = h.add_run(prefix + text)
        r.font.size = Pt(13); r.font.color.rgb = SLATE; r.bold = True
        h.paragraph_format.space_before = Pt(10); h.paragraph_format.space_after = Pt(3)
    elif level == 3:
        r = h.add_run(text)
        r.font.size = Pt(11); r.font.color.rgb = GREY; r.bold = True
        h.paragraph_format.space_before = Pt(6); h.paragraph_format.space_after = Pt(2)
    _keep(h)
    return h

def P(text, bold=False, italic=False, size=10, color=SLATE, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.15
    return p

def tech(text):
    return P(text, italic=True, size=9, color=GREY, after=3)

def B(text, prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if prefix:
        r1 = p.add_run(prefix + ' ')
        r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = SLATE
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.line_spacing = 1.1
    return p

def _img_exists(name):
    return os.path.exists(os.path.join(SS, f'{name}.png'))

def img(name, caption='', width=None):
    """Single centred image. Defaults to MOBILE_W for mobile screenshots."""
    path = os.path.join(SS, f'{name}.png')
    if not os.path.exists(path): return
    if width is None: width = MOBILE_W
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=width)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.font.size = Pt(8); r.font.color.rgb = GREY; r.italic = True
        c.paragraph_format.space_after = Pt(3)

def pair(name1, name2, cap1='', cap2='', width=None):
    """Two images side by side in a borderless table."""
    if width is None: width = PAIR_W
    p1 = os.path.join(SS, f'{name1}.png')
    p2 = os.path.join(SS, f'{name2}.png')
    if not os.path.exists(p1) and not os.path.exists(p2): return
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (path, cap) in enumerate([(p1, cap1), (p2, cap2)]):
        cell = t.cell(0, ci)
        cell.width = Inches(3.2)
        for cp in cell.paragraphs:
            cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(0)
        if os.path.exists(path):
            pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.add_run().add_picture(path, width=width)
            if cap:
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cp.add_run(cap)
                r.font.size = Pt(7.5); r.font.color.rgb = GREY; r.italic = True
                cp.paragraph_format.space_after = Pt(1)
    # remove borders
    for row in t.rows:
        for cell in row.cells:
            tc = cell._element.get_or_add_tcPr()
            tc.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0"/>'
                '<w:left w:val="none" w:sz="0" w:space="0"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                '<w:right w:val="none" w:sz="0" w:space="0"/>'
                '</w:tcBorders>'))

def stats(items):
    t = doc.add_table(rows=2, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, label) in enumerate(items):
        p1 = t.cell(0, i).paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(val); r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = BLUE
        p2 = t.cell(1, i).paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label); r2.font.size = Pt(7.5); r2.font.color.rgb = GREY
    for row in t.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC" w:val="clear"/>'))
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)

def callout(text, bg='F0F9FF'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); cell.text = ''
    cell._element.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.color.rgb = SLATE; r.italic = True
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    for pp in cell.paragraphs: pp.paragraph_format.left_indent = Inches(0.12)

def stat_line(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + '  '); r1.font.size = Pt(9.5); r1.font.color.rgb = GREY
    r2 = p.add_run(value); r2.font.size = Pt(10); r2.font.color.rgb = NAVY; r2.bold = True
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.left_indent = Inches(0.2)


# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('Chatr'); r.font.size = Pt(44); r.font.color.rgb = NAVY; r.bold = True

tp2 = doc.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('Real-Time Messaging Platform'); r2.font.size = Pt(14); r2.font.color.rgb = BLUE
tp2.paragraph_format.space_after = Pt(3)

tp3 = doc.add_paragraph(); tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run('Product Overview \u2022 Technical Architecture \u2022 Commercial Case')
r3.font.size = Pt(10.5); r3.font.color.rgb = GREY; r3.italic = True
tp3.paragraph_format.space_after = Pt(6)

tp4 = doc.add_paragraph(); tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run('A complete, production-deployed messaging platform built, tested, and deployed by a single developer.')
r4.font.size = Pt(9.5); r4.font.color.rgb = SLATE; r4.italic = True
tp4.paragraph_format.space_after = Pt(12)

img('01-landing-page', width=Inches(4.0))

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════
toc_h = doc.add_paragraph(); toc_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = toc_h.add_run('Contents'); tr.font.size = Pt(18); tr.font.color.rgb = NAVY; tr.bold = True
toc_h.paragraph_format.space_after = Pt(8)

toc = [
    ('1', 'Executive Summary', False),
    ('2', 'The Commercial Opportunity', False),
    ('3', 'Product Walkthrough', False),
    ('4', 'Messaging', False),
    ('', '4.1  Message Types', True),
    ('', '4.2  Message Actions', True),
    ('', '4.3  Real-Time Awareness', True),
    ('', '4.4  Offline & Sync', True),
    ('5', 'Groups & Social', False),
    ('6', 'The Embeddable Support Widget', False),
    ('', '6.1  How It Works', True),
    ('', '6.2  White-Label Customisation', True),
    ('', '6.3  Technical Implementation', True),
    ('7', 'AI-Powered Intelligence', False),
    ('8', 'Security, Authentication & Privacy', False),
    ('', '8.1  Authentication Methods', True),
    ('', '8.2  Privacy Controls', True),
    ('', '8.3  Rate Limiting & Protection', True),
    ('9', 'Mobile-First Design & Themes', False),
    ('10', 'Profile & Personalisation', False),
    ('11', 'Technical Architecture', False),
    ('', '11.1 The Stack', True),
    ('', '11.2 Real-Time Infrastructure', True),
    ('', '11.3 Database Design', True),
    ('', '11.4 Media Pipeline', True),
    ('', '11.5 Deployment & Scaling', True),
    ('12', 'Quality Assurance', False),
    ('13', 'Analytics Dashboard', False),
    ('14', 'Developer Experience', False),
    ('15', 'By the Numbers', False),
    ('16', 'Why Invest in Chatr', False),
]
for num, title, is_sub in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3) if is_sub else Inches(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.font.size = Pt(9) if is_sub else Pt(10)
    r.font.color.rgb = GREY if is_sub else NAVY
    r.bold = not is_sub

doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
H('Executive Summary')

P('Chatr is a fully functional, production-deployed, real-time messaging platform that demonstrates the breadth of a funded engineering team \u2014 delivered by a single developer in 22 days. It is not a prototype or proof of concept. It is a working product with 50+ user-facing features, 1,300+ automated tests, and a deployment running on AWS infrastructure.', bold=True)

P('For a commercial audience, Chatr shows what a complete product looks like when messaging, AI, and customer support converge into a single platform. Its embeddable chat widget allows any business to add real-time customer support to their website with a single line of code \u2014 competing directly with tools like Intercom ($39\u2013$99/seat/month) at zero recurring cost.')

P('For a technical audience, Chatr demonstrates mastery across frontend development (React 19, Next.js 16), backend engineering (Node.js, Express, Socket.IO), database design (PostgreSQL, Prisma), caching infrastructure (Redis), AI integration (OpenAI GPT-4o-mini), cloud deployment (AWS), and automated testing (Jest, Playwright). Every layer is production-grade, documented, and covered by automated tests.')

stats([('50+', 'User Features'), ('1,300+', 'Automated Tests'), ('78,000+', 'Lines of Code'), ('22', 'Days to Ship')])

callout('This document serves as both a commercial presentation and a technical reference. Each section explains what a feature does, why it matters commercially, and how it is implemented technically.')


# ═════════════════════════════════════════════════════════════════════════════
# 2. THE COMMERCIAL OPPORTUNITY
# ═════════════════════════════════════════════════════════════════════════════
H('The Commercial Opportunity')

H('The Market Problem', level=2)
P('Every business needs real-time communication. Internal teams need to collaborate instantly. Support teams need to respond to customers while they are still on the website. Users expect the experience they get from consumer apps \u2014 instant delivery, typing indicators, read receipts, voice notes, file sharing, and a mobile-first interface that works flawlessly on any device.')

P('Today, companies face a costly trade-off:')
B('Buy a third-party SaaS solution (Intercom, Zendesk Chat, Drift) and pay $39\u2013$99 per agent per month, with limited customisation, vendor lock-in, and no control over user data.', prefix='Option A:')
B('Build from scratch, spending 3\u20136 months on WebSocket infrastructure, message queuing, presence tracking, and delivery guarantees before writing a single user-facing feature.', prefix='Option B:')

H('The Chatr Answer', level=2)
P('Chatr eliminates that trade-off. It is a fully functional messaging platform that delivers enterprise-grade features at zero licensing cost. It is open, extensible, and built on industry-standard technology (React, Node.js, PostgreSQL, Redis) that any JavaScript developer can understand, modify, and extend on day one.', bold=True)

P('Its embeddable chat widget turns any website into a live support channel \u2014 creating a direct, zero-friction communication bridge between businesses and their customers. That is where the commercial value lives.')

P('Live chat is one of the fastest-growing segments in customer support. Businesses using real-time chat see 48% higher revenue per chat hour compared to email support, and customers report 73% satisfaction rates with live chat \u2014 the highest of any support channel.')

P('The market for live chat software is dominated by expensive SaaS products:')
B('Intercom: $39\u2013$99/seat/month, with feature gating and usage limits.')
B('Drift: $50\u2013$150/seat/month, focused on sales automation.')
B('Zendesk Chat: $19\u2013$99/seat/month, bundled with ticketing overhead.')
B('Chatr Widget: $0/seat, full ownership, no recurring cost, no vendor lock-in.', prefix='\u2192')

P('For a company with a 10-person support team, that is $4,700\u2013$11,900/year in savings compared to Intercom alone. Beyond cost savings, Chatr gives businesses complete control over their data, full white-label customisation, and the ability to extend functionality without waiting for a vendor\u2019s roadmap.', bold=True)


# ═════════════════════════════════════════════════════════════════════════════
# 3. PRODUCT WALKTHROUGH
# ═════════════════════════════════════════════════════════════════════════════
H('Product Walkthrough')

P('The screenshot below shows the main conversation screen. This single view demonstrates over a dozen capabilities working together in real time: groups, direct messages, AI conversations, guest visitors from the widget, typing indicators, unread badges, online presence, friend badges, and AI-generated conversation summaries.')

img('03-conversations', 'Main conversation list \u2014 groups, DMs, AI bot, widget guests, typing indicators, presence, unread badges, and smart summaries')

P('Every row in the conversation list is information-dense by design. At a glance, users can see:')
B('Whether the contact is online (green dot), away (amber dot), or offline (grey dot) via real-time presence tracking.', prefix='Presence \u2014')
B('If someone is currently typing, the last-message preview is replaced with an animated "typing\u2026" indicator, visible without opening the conversation.', prefix='Typing indicators \u2014')
B('Per-conversation unread message counts displayed as badges on each row, plus an aggregate badge on the bottom navigation tab for total unread messages across all conversations.', prefix='Unread badges \u2014')
B('"Friend", "Group", "AI", and "Guest" badges distinguish conversation types at a glance. These trust indicators help users prioritise which conversations to open first.', prefix='Conversation badges \u2014')
B('When enabled, the system generates AI summaries that replace the last-message preview with a concise description of what was discussed, animated with a flip transition.', prefix='Smart summaries \u2014')
B('Messages from unknown contacts are separated into a "Requests" tab. Each request can be accepted, declined, or blocked with a single tap.', prefix='Message requests \u2014')

P('The search bar at the top filters conversations in real time as you type, matching against contact names, group names, and message content. The bottom navigation provides instant access to Chats, Friends, Groups, and your Profile.')

H('Registration & Login', level=2)

P('New users register with a username, email address, and password. During registration, a real-time password strength indicator shows the strength of the chosen password as the user types: Weak (red), Fair (amber), Good (green), Strong (bright green). The strength meter evaluates length, character variety (uppercase, lowercase, numbers, symbols), and common patterns.')

pair('30-register-form', '02-login', 'Registration with password strength', 'Login screen')

P('After submitting the registration form, a 6-digit one-time code is sent to the user\u2019s email address. The user must enter this code to verify their identity and activate their account. This prevents fake account creation and ensures every user has a valid email address.')

P('Authentication supports four methods: email/password with verification code, SMS verification, email-based login codes, and optional TOTP two-factor authentication via any authenticator app (Google Authenticator, Authy, Microsoft Authenticator).')
tech('JWT access tokens (localStorage) with long-lived HttpOnly refresh cookies. Expired/revoked tokens blacklisted in Redis. Rate-limited login attempts prevent brute-force attacks. Automatic token refresh keeps users logged in without re-authenticating.')


# ═════════════════════════════════════════════════════════════════════════════
# 4. MESSAGING
# ═════════════════════════════════════════════════════════════════════════════
H('Messaging')

P('Chatr delivers a messaging experience that matches or exceeds what users expect from WhatsApp, iMessage, and Slack. Every message type, interaction pattern, and real-time indicator found in those apps has been implemented, tested, and refined.')

P('This section provides a detailed breakdown of every messaging capability, how it works from the user\u2019s perspective, and the technology powering it behind the scenes.')

pair('04-chat-view', '04b-chat-view-top',
     'Code blocks, reactions, replies, edits, receipts',
     'Voice waveforms, shared media, link previews')

# ── 4.1 Message Types ────────────────────────────────────────────────────
H('Message Types', level=2)

P('Chatr supports seven distinct message types. Each type has its own rendering, interaction model, and metadata. Users never need to think about "types" \u2014 they simply send content and Chatr handles everything.')

H('Text Messages', level=3)
P('The foundation of any messaging platform. Text messages in Chatr include delivery status tracking through three states: "sending" (clock icon), "delivered" (single tick), and "read" (double tick, blue). Each message displays a timestamp, and messages are grouped by date with separator headers. Consecutive messages from the same sender are visually grouped to reduce clutter.')
tech('Transmitted via WebSocket (Socket.IO) for instant delivery. When the recipient is online, the message arrives in under 100ms. If the recipient is offline, the message is stored in PostgreSQL and delivered when they reconnect. Delivery acknowledgements flow back through the same WebSocket channel, updating the sender\u2019s UI in real time.')

H('Voice Messages', level=3)
P('Users record voice notes directly from the chat input bar. During recording, the other participant sees a "recording\u2026" indicator in real time. Once sent, voice messages display an interactive waveform visualisation showing the audio shape, a duration counter, and playback controls (play, pause, scrub). When the recipient plays the message, the sender receives a "listened" receipt \u2014 equivalent to a read receipt, but specific to audio.')
tech('Recorded using the Web Audio API and MediaRecorder, encoded as WebM/Opus, uploaded to AWS S3, and streamed on playback. The waveform is pre-computed during recording by sampling audio amplitude data. Played voice notes are cached locally in IndexedDB so repeat listens are instant, even offline.')

H('Image Sharing', level=3)
P('Images appear as inline previews within the conversation, proportionally scaled to fit the message bubble. Tapping any image opens a fullscreen lightbox with a dimmed background, allowing users to view the image at its original resolution. The lightbox supports pinch-to-zoom on mobile and close-on-escape on desktop.')
img('32-image-lightbox', 'Fullscreen image lightbox')
tech('On upload, images are processed server-side using the Sharp library to generate multiple resolution variants (thumbnail, medium, full). The frontend loads the smallest variant for the inline preview and the full resolution only when the lightbox is opened, minimising bandwidth usage. All images are stored on S3 with unique hashed filenames.')

H('Video Sharing', level=3)
P('Videos are uploaded and displayed as inline players with a thumbnail preview, file name, and duration. Users tap to play directly within the conversation without leaving the chat. The video player supports standard controls: play, pause, seek, volume, and fullscreen.')

H('File Attachments', level=3)
P('Chatr supports file uploads up to 50 MB, covering PDFs, Word documents, Excel spreadsheets, ZIP archives, and more. Each attachment is displayed with a type-specific icon (determined by MIME type), the file name, file size, and a download link. Files are uploaded to S3 and streamed directly to the recipient.')

H('Link Previews', level=3)
P('When a message contains a URL, Chatr automatically fetches the page\u2019s Open Graph metadata (title, description, image, favicon) and renders a rich preview card below the message text. Preview cards show the site name, page title, a description excerpt, and a thumbnail image \u2014 the same way link previews function in Slack and iMessage.')
tech('The preview is generated server-side: when the backend detects a URL in an outgoing message, it fetches the target page, parses the og:title, og:description, og:image, and favicon meta tags, and stores the preview metadata alongside the message. The frontend then renders a clickable card that opens the link in a new tab.')

H('Code Blocks', level=3)
P('Messages containing code fenced in triple backticks are rendered with syntax highlighting, automatic language detection, and a "Copy" button that copies the code to the clipboard with one click. The code block supports over 40 programming languages and is styled with a dark monospace theme for readability. This feature is particularly useful for developer-to-developer communication and technical support scenarios.')
img('41-code-block', 'Syntax-highlighted TypeScript code block with Copy button')

# ── 4.2 Message Actions ──────────────────────────────────────────────────
H('Message Actions', level=2)

P('Every message in Chatr is interactive. Users can react, reply, edit, and unsend messages, giving them the same control they have in modern consumer messaging apps.')

H('Emoji Reactions', level=3)
P('Users react to any message by tapping and holding (mobile) or hovering and clicking the reaction icon (desktop). The reaction bar offers six preset emoji. Once applied, reactions appear as small badges below the message bubble, showing the emoji icon and a count of how many people reacted with it. Hovering over a reaction badge reveals a tooltip listing each person who reacted. Users can remove their own reaction by tapping it again.')
tech('Reactions are stored as separate database records linked to the message ID, enabling multiple reactions per message from different users. Reaction additions and removals are broadcast in real time via Socket.IO, so all participants see updates instantly.')

H('Reply to Message', level=3)
P('Users can reply to any message in the conversation by swiping right (mobile) or selecting "Reply" from the context menu (desktop). The reply creates a quoted preview above the new message, showing the original sender\u2019s name, the content type (text, voice, image, etc.), and a truncated preview of the content. Tapping the quoted preview scrolls the conversation to the original message and briefly highlights it, making it easy to follow conversation threads.')

pair('42-reactions', '43-reply-thread',
     'Reaction badges with emoji counts',
     'Reply with quoted original message')

H('Edit Sent Messages', level=3)
P('Users can edit any message they have sent. On desktop, pressing the Up arrow in an empty input field automatically enters edit mode for the most recent message \u2014 the same shortcut used in Slack. Alternatively, users can right-click (or long-press on mobile) any of their own messages and select "Edit" from the context menu. Edited messages display a small "(edited)" label next to the timestamp.')
tech('Edits are versioned: the backend stores a complete edit history (original content, new content, timestamp) for audit purposes. The edit event is broadcast via Socket.IO, and all participants see the updated content in real time.')

H('Unsend Messages', level=3)
P('Users can unsend (delete) any message they have sent. The message is removed from the conversation for all participants \u2014 both sender and recipient(s) in DMs and groups. A small "message was deleted" placeholder optionally appears in the conversation flow. This is equivalent to the "Delete for Everyone" feature in WhatsApp.')

H('Emoji Picker', level=3)
P('The message input includes a dedicated emoji button that opens a full emoji picker panel. The picker includes category tabs (Smileys, People, Animals, Food, Travel, Objects, Symbols, Flags), a search bar for finding specific emoji by name, and a "Recently Used" section that remembers the user\u2019s most-used emoji. Selecting an emoji inserts it at the cursor position in the message input.')
img('22-emoji-picker', 'Full emoji picker with categories, search, and recently-used')

# ── 4.3 Real-Time Awareness ──────────────────────────────────────────────
H('Real-Time Awareness', level=2)

P('One of Chatr\u2019s most distinctive qualities is the density of its real-time feedback. Users always know what is happening \u2014 who is online, who is typing, who has read their message, and even what the other person is writing, character by character.')

H('Typing Indicators in Chat', level=3)
P('When the other participant begins typing in a direct message, an animated "typing\u2026" indicator with bouncing dots appears at the bottom of the conversation, just above the input bar. The indicator appears within 200ms of the first keystroke and disappears 3 seconds after the last keystroke (or immediately when the message is sent). In group chats, the indicator shows the name of the person typing.')

H('Typing Indicators on the Chat List', level=3)
P('This is a feature that most messaging apps do not offer. When someone is typing a message to you, the last-message preview on the conversation list is replaced with a "typing\u2026" label. This means users can see who is composing a message without opening the conversation \u2014 saving time and providing ambient awareness of activity across all their chats.')

pair('26-typing-in-chat', '25-typing-chat-list',
     'Typing indicator in chat (animated dots)',
     '"typing\u2026" replacing last-message on list')

tech('Typing events are emitted via Socket.IO with debouncing: the frontend sends a typing:start event on the first keystroke, then suppresses additional events for 2 seconds. The backend relays the event to the recipient\u2019s Socket.IO room. The 3-second timeout is managed client-side.')

H('Ghost Typing', level=3)
P('Ghost typing is an optional mode (toggled in Settings) that streams every character the other person types in real time, letter by letter. Instead of seeing a generic "typing\u2026" indicator, you see the actual text appearing character by character as the other person composes their message. This creates an intimate, almost telepathic communication experience that no mainstream messaging app offers.')
tech('Ghost typing works by emitting a typing:ghost event on every keystroke, containing the current content of the input field. The recipient\u2019s client renders the content in a preview bubble that updates with each incoming event. To prevent excessive bandwidth, events are batched into 50ms windows.')

H('Audio Recording Indicator', level=3)
P('When a user begins recording a voice note, the other participant sees a "recording\u2026" indicator in the chat view, letting them know a voice message is on its way. This prevents the awkward silence that occurs when someone is recording audio but the other person does not know they are still engaged.')

H('Online Presence', level=3)
P('Every avatar in Chatr \u2014 on the conversation list, in chat headers, in group member lists, and on profile pages \u2014 displays a coloured presence dot. Green means online, amber means away (idle for 5+ minutes), and grey means offline. Below the avatar, offline contacts show a "last seen X ago" timestamp so users know how recently the person was active.')
tech('Presence is tracked via Redis. When a user connects (or their Socket.IO session heartbeats), their status is set to "online" in Redis with a TTL. When they disconnect or idle, the status transitions to "away" and then "offline". Presence changes are broadcast to all contacts via Socket.IO.')

H('Read Receipts', level=3)
P('Every text message transitions through three delivery states, each with a distinct icon: sending (clock), delivered (single tick), and read (double tick in blue). When the recipient opens the conversation and the message scrolls into view, a "read" event is emitted back to the sender, and the icon updates in real time. Voice messages have an additional "listened" state.')

# ── 4.4 Offline & Sync ───────────────────────────────────────────────────
H('Offline & Sync', level=2)

P('Chatr is designed to work reliably even when the network is unavailable. This matters for mobile users in areas with poor connectivity, for desktop users switching between networks, and for any scenario where uninterrupted communication is critical.')

H('IndexedDB Cache', level=3)
P('All conversations, contacts, messages, and metadata are cached locally in the browser\u2019s IndexedDB. When the app loads, it renders immediately from the local cache while syncing with the server in the background. This means the app feels instant \u2014 no loading spinner, no blank screen, no waiting for API responses.')

H('Outbound Message Queue', level=3)
P('If a user sends a message while offline, the message is added to a local outbound queue. The chat UI shows the message immediately (with a "sending" clock icon), and a small badge on the input bar shows the count of queued messages. When the connection is restored, queued messages are sent automatically in order, and the delivery icons update as acknowledgements arrive.')

H('Audio Cache', level=3)
P('Voice messages are cached locally after the first playback. Subsequent listens play instantly from the local cache, even if the device is offline. This is particularly useful for users who revisit voice notes frequently.')

H('Storage Management', level=3)
P('Users can view a breakdown of how much local storage each cache category uses (messages, voice notes, audio, media) in Settings. A visual bar chart shows the proportion of space used by each category. Users can reset individual caches or clear all local data with a single tap.')


# ═════════════════════════════════════════════════════════════════════════════
# 5. GROUPS & SOCIAL
# ═════════════════════════════════════════════════════════════════════════════
H('Groups & Social')

P('Team collaboration requires structure. Chatr provides full-featured group chats with role management, invite controls, and a social layer that gives users complete control over who can reach them.')

pair('06-groups', '05-friends',
     'Groups with member counts and search',
     'Friends with online presence')

H('Group Chat', level=2)
P('Groups in Chatr are designed for team collaboration. Each group has a name, an optional avatar image, a cover image, and a description \u2014 essentially a profile page for the group. All seven message types (text, voice, image, video, file, code, links) are supported in group conversations, along with reactions, replies, and edits.')

pair('23-group-chat', '34-group-members',
     'Group conversation',
     'Member list with role badges')

H('Role Management', level=3)
P('Groups use a three-tier role hierarchy: Owner, Admin, and Member. Each role has different permissions:')
B('Can promote members to Admin, demote Admins to Member, transfer ownership to another member, remove any member, delete the group, and edit group details (name, avatar, cover, description).', prefix='Owner \u2014')
B('Can add new members, remove Members (but not other Admins or the Owner), and edit group details.', prefix='Admin \u2014')
B('Can send messages, react, reply, and leave the group.', prefix='Member \u2014')
P('Role badges (Owner, Admin) are displayed next to member names in the group member list, making the hierarchy immediately visible.')

H('Invitations', level=3)
P('Owners and Admins can invite new members by searching across all Chatr users from a dedicated "Add Members" panel. Invited users receive a notification and can accept or decline the invitation. Pending invitations appear as a count badge on the Groups tab. Declined invitations are silently removed.')

H('Friends & Social Layer', level=2)
P('The friend system controls who can message you directly and provides visual trust indicators across the interface.')

H('Friend Requests', level=3)
P('Users can send friend requests to any Chatr user found via search. The recipient sees the request in a dedicated "Incoming" panel and can accept, decline, or block the sender. Accepted contacts appear with a "Friend" badge on the conversation list. The friend list shows all accepted contacts with their current online status and supports search filtering.')

H('Blocking', level=3)
P('Blocking is comprehensive: a blocked user cannot search for you, send you messages, view your profile, see your online status, or send you friend requests. From the blocked user\u2019s perspective, you effectively disappear from the platform. Blocked users are managed from a dedicated panel in Settings, where they can be unblocked at any time.')

H('User Search', level=3)
P('A global search bar allows users to find any Chatr user by name or username. Search results show the user\u2019s avatar, display name, and current friend status. From search results, users can start a new conversation, send a friend request, or view the user\u2019s profile.')


# ═════════════════════════════════════════════════════════════════════════════
# 6. THE EMBEDDABLE WIDGET
# ═════════════════════════════════════════════════════════════════════════════
H('The Embeddable Support Widget')

callout('The widget is the feature that transforms Chatr from a messaging app into a revenue-generating customer support platform. This is where the commercial value proposition is strongest.')

H('How It Works', level=2)
P('Any website can add live customer support by pasting a single line of JavaScript into their HTML. A floating chat bubble appears in the bottom-right corner of the page. When a visitor clicks the bubble, a chat panel opens with a customisable greeting, asking for their name and initial question. No account creation, no email address, no sign-up flow \u2014 zero friction.')

P('Once the visitor submits their name and message, a real-time connection is established via Socket.IO. The message appears instantly in the support agent\u2019s Chatr inbox, tagged with a "Guest" badge so the agent can distinguish customer enquiries from internal conversations. The agent replies directly from Chatr, and the visitor sees the response in real time on the website. The conversation persists for 24 hours, so visitors can close the tab and return later to continue the conversation.')

pair('11-widget-intro', '11b-widget-form-filled',
     'Widget greeting panel',
     'Visitor fills in name and question')

pair('11c-widget-chat', '11d-widget-conversation',
     'Message arrives in agent\u2019s inbox',
     'Live two-way conversation')

H('Why This Matters Commercially', level=2, numbered=False)
P('The market for live chat software is dominated by expensive SaaS products. Chatr delivers the same core functionality \u2014 real-time messaging, file sharing, persistent sessions, agent routing \u2014 with full code ownership. A business deploys it once and pays nothing ongoing.', bold=True)

H('White-Label Customisation', level=2)
P('The widget is fully white-labelled. A visual "Palette Designer" in the Chatr dashboard lets agents configure:')
B('Primary colour, background colour, text colour, and header colour via colour pickers.')
B('Dark mode or light mode, toggled with a single switch.')
B('Custom greeting text and placeholder messages.')
B('Preset colour themes for quick configuration.')
B('A one-click "Copy Embed Code" button that generates the complete HTML snippet.')

img('37-widget-palette-designer', 'Widget Palette Designer \u2014 colours, themes, and embed code', width=WIDE_W)

H('Technical Implementation', level=2)
P('The widget is a standalone JavaScript file (chatr.js) that can be loaded from any domain. It injects its own DOM elements, styles, and event handlers without interfering with the host page. The widget creates a guest session on the backend (no authentication required), establishes a Socket.IO WebSocket connection, and routes messages to the designated agent. Sessions are stored in localStorage with a 24-hour TTL.')
P('The widget supports all DM features: text messages, voice notes, file attachments, link previews, typing indicators, and read receipts. It is not a stripped-down version \u2014 it is the full Chatr messaging experience embedded in a third-party website.')


# ═════════════════════════════════════════════════════════════════════════════
# 7. AI-POWERED INTELLIGENCE
# ═════════════════════════════════════════════════════════════════════════════
H('AI-Powered Intelligence')

P('Chatr integrates AI at two levels: a conversational assistant that users interact with directly, and a background intelligence layer that processes conversations automatically.')

H('Luna \u2014 AI Chat Assistant', level=2)
P('Luna is Chatr\u2019s built-in AI chatbot, powered by OpenAI\u2019s GPT-4o-mini model. She appears as a regular contact in the conversation list with a distinctive teal ring on her avatar. Users can ask Luna questions, request help, have casual conversations, or use her as a knowledge base \u2014 all without leaving the messaging interface.')

P('Unlike most AI integrations that feel bolted on, Luna is a first-class citizen in Chatr. Her conversations support the same features as human conversations: typing indicators (Luna "types" while generating a response), message history, and the full chat UI. This means there is zero learning curve \u2014 users interact with Luna exactly the same way they interact with anyone else.')

img('20-luna-chat', 'Luna AI assistant providing detailed, contextual responses')

tech('Luna\u2019s responses are generated by sending the conversation history (up to a configurable context window) to the OpenAI API with a system prompt that establishes her personality: helpful, concise, and knowledgeable. Responses stream back token by token, which is why the typing indicator appears while Luna is "thinking". The backend handles rate limiting and error recovery to ensure Luna is always available.')

H('Automatic Conversation Summaries', level=2)
P('When a conversation accumulates enough messages, the system automatically generates a concise AI summary that describes the key topics discussed. These summaries appear on the conversation list, replacing the last-message preview with a one-line summary of the entire conversation. Summaries animate into view with a smooth flip transition, so users can scan the state of all their conversations without opening any of them.')
P('This is particularly valuable for teams: a manager can glance at the conversation list and immediately understand what each thread is about, without reading through the messages.')

H('Toast Notifications', level=2)
P('Incoming messages trigger in-app toast notifications that appear at the top of the screen with the sender\u2019s avatar, name, and a preview of the message content. Toast notifications are non-intrusive (they auto-dismiss after 5 seconds) but clickable \u2014 tapping a toast opens the relevant conversation. This ensures users never miss a message, even when viewing a different conversation or another page in the app.')


# ═════════════════════════════════════════════════════════════════════════════
# 8. SECURITY, AUTHENTICATION & PRIVACY
# ═════════════════════════════════════════════════════════════════════════════
H('Security, Authentication & Privacy')

P('Chatr implements enterprise-grade identity verification and granular privacy controls. This section covers every authentication method, the privacy settings available to users, and the security measures protecting the platform.')

H('Authentication Methods', level=2)

H('Registration', level=3)
P('New users register with a username, email address, and password. During registration, a real-time password strength indicator shows the strength of the chosen password as the user types: Weak (red), Fair (amber), Good (green), Strong (bright green). The strength meter evaluates length, character variety (uppercase, lowercase, numbers, symbols), and common patterns.')
P('After submitting the form, a 6-digit one-time code is sent to the user\u2019s email address. The user must enter this code to verify their identity and activate their account.')

H('Login', level=3)
P('Users log in with their email/username and password. After successful credential verification, an additional login verification code is sent via email or SMS (depending on user preference). This two-step login process ensures that even if a password is compromised, the account remains secure.')
tech('Authentication is implemented using JSON Web Tokens (JWT). On successful login, the server issues a short-lived access token (stored in localStorage) and a long-lived refresh token (stored in an HttpOnly cookie). When the access token expires, the client automatically refreshes it using the refresh token without requiring the user to log in again. Expired and revoked tokens are blacklisted in Redis.')

H('Two-Factor Authentication (2FA)', level=3)
P('Users can enable optional two-factor authentication from Settings. The setup flow generates a QR code that the user scans with any TOTP-compatible authenticator app (Google Authenticator, Authy, Microsoft Authenticator, etc.). Once enabled, every login requires both the password and a 6-digit TOTP code from the authenticator app. Backup codes are provided during setup in case the user loses access to their authenticator.')
tech('2FA is implemented using the TOTP standard (RFC 6238). The server generates a unique secret key per user, which is encoded in the QR code. On each login attempt, the server validates the provided code against the current time window. Invalid codes are rate-limited to prevent brute-force attacks.')

H('SMS Verification', level=3)
P('Users can add and verify a phone number. The backend sends a one-time SMS code via a gateway service, and the user enters the code in the app. Verified phone numbers can be used as a secondary login verification channel.')

H('Password Recovery', level=3)
P('A "Forgot Password" link on the login screen initiates a secure email-based reset flow. The user enters their email, receives a one-time reset link, clicks it, and sets a new password. The reset link expires after a configurable time window and can only be used once.')

H('Privacy Controls', level=2)
P('Chatr gives users granular control over their personal information visibility. From the Privacy panel in Settings, users can configure who can see each piece of their profile information:')
B('Online status (the green/amber/grey presence dot): visible to Everyone, Friends only, or Nobody.')
B('Full name: visible to Everyone, Friends only, or Nobody.')
B('Phone number: visible to Everyone, Friends only, or Nobody.')
B('Email address: visible to Everyone, Friends only, or Nobody.')
B('Gender: visible to Everyone, Friends only, or Nobody.')
B('Join date: visible to Everyone, Friends only, or Nobody.')
img('24-privacy-settings', 'Privacy settings \u2014 per-field visibility controls')
tech('These settings are enforced server-side: when another user requests a profile, the backend checks the visibility settings for each field and strips any fields the requesting user is not permitted to see. This means the privacy controls cannot be bypassed by inspecting network requests.')

H('Rate Limiting & Protection', level=2)
P('All sensitive endpoints (login, registration, password reset, verification codes, message sending) are rate-limited using Redis-backed sliding window counters. A user who submits too many failed login attempts is temporarily locked out. A client sending too many messages per second is throttled. This protects against brute-force attacks, credential stuffing, and denial-of-service attempts.')


# ═════════════════════════════════════════════════════════════════════════════
# 9. MOBILE-FIRST DESIGN & THEMES
# ═════════════════════════════════════════════════════════════════════════════
H('Mobile-First Design & Themes')

P('Every screenshot in this document was captured at mobile viewport resolution (390\u00d7844, iPhone 14). This is deliberate: Chatr is designed for mobile first and scales up to desktop, not the other way around. The interface uses bottom tab navigation (Chats, Friends, Groups, Settings), touch-optimised tap targets, swipe gestures for message actions, and iOS safe-area insets to avoid the notch and home indicator.')

P('On desktop and tablet, the layout adapts to use the additional screen space: the conversation list appears as a persistent sidebar alongside the chat view, eliminating the need to navigate back and forth. The responsive breakpoint is seamless \u2014 users never see a "desktop version" or a "mobile version", just the same app adapting intelligently to the available space.')

H('Dark & Light Themes', level=2)
P('Users can switch between dark and light themes with a single tap in Settings. The theme preference is saved locally and applied instantly \u2014 no page reload, no flicker. The dark theme uses a deep navy palette with blue accents, optimised for OLED screens (true black backgrounds reduce power consumption). The light theme uses a clean white palette with subtle grey borders and blue highlights.')

pair('18-dark-theme', '27-dark-theme-chat',
     'Dark theme \u2014 conversation list',
     'Dark theme \u2014 chat view')

pair('19-light-theme', '36-light-theme-chat',
     'Light theme \u2014 conversation list',
     'Light theme \u2014 chat view')

tech('The theme system is implemented using CSS custom properties (variables) that are toggled at the document root. Every colour in the application is defined as a variable, so the entire visual identity changes with a single class swap. The theme preference is stored in localStorage and applied before the first paint via a blocking script, preventing the "flash of wrong theme" that many apps suffer from.')


# ═════════════════════════════════════════════════════════════════════════════
# 10. PROFILE & PERSONALISATION
# ═════════════════════════════════════════════════════════════════════════════
H('Profile & Personalisation')

P('Users own their identity on Chatr. The profile and settings system gives them tools to express themselves and control their experience.')

H('Profile', level=2)
P('Each user has a profile page with a circular avatar photo and a 16:9 cover banner image. Both images can be uploaded and cropped directly in the app using built-in circular and rectangular crop tools. The profile also displays the user\u2019s display name, bio, username, and any personal details they have chosen to make visible.')

H('Settings', level=2)
P('The Settings screen is a comprehensive control panel:')
B('Theme toggle (dark/light) with instant preview.')
B('Ghost typing toggle (enable/disable character-by-character typing preview).')
B('Privacy settings (per-field visibility controls, as detailed in Section 8).')
B('Blocked users management panel.')
B('Storage usage chart with per-category breakdown and one-tap reset.')
B('Account settings (display name, bio, email, phone number).')
B('Two-factor authentication setup.')
B('Logout.')

pair('08-profile', '07-settings',
     'User profile with avatar and cover',
     'Settings panel')

img('35-settings-storage', 'Storage usage breakdown \u2014 per-category cache sizes')


# ═════════════════════════════════════════════════════════════════════════════
# 11. TECHNICAL ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
H('Technical Architecture')

P('This section is for technical evaluators who need to understand how Chatr is built, what technologies it uses, and why each choice was made. The architecture follows industry best practices and is designed for production reliability, horizontal scalability, and developer productivity.')

H('The Stack', level=2)

P('Frontend \u2014 Next.js 16, React 19, TypeScript.', bold=True)
P('The frontend is a Next.js 16 application using React 19 with the App Router. All code is TypeScript with strict mode enabled. State management uses React Context with custom hooks (no Redux). Animations are powered by Framer Motion. The chat interface uses Socket.IO for real-time communication and IndexedDB (via a custom wrapper) for offline data caching.')

P('Backend \u2014 Node.js, Express, TypeScript.', bold=True)
P('The backend is a Node.js/Express server, also in strict TypeScript. It serves both a REST API (70+ endpoints) and a WebSocket server (40+ event types) via Socket.IO. The backend handles authentication, message routing, file uploads, AI integration, email/SMS delivery, and all business logic.')

P('Database \u2014 PostgreSQL 16, Prisma ORM.', bold=True)
P('PostgreSQL 16 is the primary data store, accessed through Prisma ORM for type-safe queries with automatic migrations. The schema includes 9 models: User, Conversation, Message, Group, GroupMember, Friendship, Reaction, MessageEditHistory, and Session.')

P('Caching & Pub/Sub \u2014 Redis 7.', bold=True)
P('Redis serves four critical functions: (1) presence tracking (online/away/offline status with TTLs), (2) rate limiting (sliding window counters for sensitive endpoints), (3) token blacklisting (revoked JWTs stored until their natural expiry), and (4) cross-instance pub/sub via the Socket.IO Redis adapter, enabling horizontal scaling.')

P('AI \u2014 OpenAI GPT-4o-mini.', bold=True)
P('The chatbot (Luna) and automatic conversation summaries use OpenAI\u2019s GPT-4o-mini model via the OpenAI Node.js SDK. Prompts are carefully engineered for conciseness and helpfulness.')

P('Storage \u2014 AWS S3.', bold=True)
P('All user-uploaded media (images, videos, voice notes, files, avatars, cover images) is stored in S3. Images are processed server-side using the Sharp library, which generates multiple resolution variants (thumbnail, medium, original) on upload.')

H('Real-Time Infrastructure', level=2)

P('Real-time communication is the heart of Chatr. Every message, typing indicator, presence update, reaction, read receipt, and notification is delivered via WebSockets.')

H('Message Delivery Pipeline', level=3)
P('When User A sends a message to User B, the following sequence occurs:')
B('User A\u2019s frontend emits a "message:send" event via Socket.IO, containing the message content, type, and recipient ID.')
B('The backend receives the event, validates the payload, and persists the message to PostgreSQL.')
B('The backend emits a "message:new" event to User B\u2019s Socket.IO room. If User B has multiple devices connected, all of them receive the message simultaneously.')
B('The backend emits a "message:delivered" acknowledgement back to User A, updating the delivery status from "sending" to "delivered".')
B('When User B opens the conversation and the message scrolls into view, User B\u2019s frontend emits a "message:read" event. The backend relays this to User A, updating the status to "read".')
P('The entire pipeline from send to receive completes in under 100ms on a typical connection. If User B is offline, the message is stored and delivered on their next connection.')

H('Horizontal Scaling', level=3)
P('Chatr uses the Socket.IO Redis adapter, which means multiple backend instances can run behind a load balancer and still deliver messages correctly. When Instance 1 receives a message for a user connected to Instance 2, the Redis pub/sub layer ensures the message is relayed to Instance 2 transparently. Scaling horizontally means adding more EC2 instances behind the load balancer \u2014 no code changes required.')

H('Database Design', level=2)

P('The PostgreSQL database uses 9 models managed by Prisma ORM:')
B('Stores credentials, profile data, settings, presence, and guest/widget sessions.', prefix='User \u2014')
B('Represents a DM thread between two users. Stores last message timestamp for sorting.', prefix='Conversation \u2014')
B('Contains the message content, type, sender, recipient/group, delivery status, reply reference, and timestamps. Indexed on sender, recipient, group, and created date.', prefix='Message \u2014')
B('Group name, description, avatar, cover image, and creation metadata.', prefix='Group \u2014')
B('Junction table linking users to groups with their role (Owner, Admin, Member) and invite status.', prefix='GroupMember \u2014')
B('Represents a friend connection between two users with status (pending, accepted, blocked).', prefix='Friendship \u2014')
B('Links an emoji to a message and user, with a unique constraint preventing duplicate reactions.', prefix='Reaction \u2014')
B('Stores the full edit history for a message (original and updated content, timestamp).', prefix='MessageEditHistory \u2014')
B('Tracks active sessions for token management and multi-device support.', prefix='Session \u2014')

H('Media Pipeline', level=2)

P('When a user uploads a file, the following happens:')
B('The frontend sends the file to the backend via a multipart POST request.')
B('The backend validates the file type and size (max 50 MB).')
B('If the file is an image, Sharp generates three variants: a 200px thumbnail, a 600px medium preview, and the original. All three are uploaded to S3.')
B('If the file is not an image (video, PDF, archive), it is uploaded to S3 as-is.')
B('The backend returns the S3 URLs, which are stored in the message record and delivered to the recipient.')
B('The frontend renders the appropriate preview (image thumbnail, video player, file icon with download link) based on the file type.')

H('Deployment & Scaling', level=2)

P('The production deployment uses five AWS services:')
B('The Node.js backend runs in PM2 cluster mode across all available CPU cores on an EC2 instance.', prefix='EC2 \u2014')
B('Managed PostgreSQL with automatic backups, failover, and encryption at rest.', prefix='RDS \u2014')
B('Managed Redis for presence, rate limiting, pub/sub, and token blacklisting.', prefix='ElastiCache \u2014')
B('Media storage with server-side encryption and lifecycle policies.', prefix='S3 \u2014')
B('Reverse proxy, TLS termination, WebSocket upgrade handling, and static file serving.', prefix='Nginx \u2014')

P('The same stack can be run locally using Docker Compose, which provisions PostgreSQL, Redis, the backend, and the frontend in a single command. Development and production environments are architecturally identical.')


# ═════════════════════════════════════════════════════════════════════════════
# 12. QUALITY ASSURANCE
# ═════════════════════════════════════════════════════════════════════════════
H('Quality Assurance')

P('Chatr has the test coverage of a funded engineering team. Over 1,300 automated tests ensure that every feature works, every edge case is handled, and every deployment is safe.')

stats([('1,300+', 'Total Tests'), ('855', 'Frontend'), ('305', 'Backend'), ('156', 'End-to-End')])

H('Test Strategy', level=2)
P('The test suite is structured in three tiers, following the testing pyramid:')

H('Frontend Unit & Integration Tests (855)', level=3)
P('Every React component, custom hook, context provider, form, dialog, and page in the frontend has corresponding tests. These tests verify rendering, user interactions, state changes, error handling, and accessibility (ARIA attributes, keyboard navigation, screen reader announcements). The frontend test suite runs in under 30 seconds and covers 99% of the frontend codebase.')
P('Key areas covered: all messaging UI (send, receive, edit, unsend, reactions, replies), conversation list rendering and filtering, group management flows, friend request flows, settings toggles, privacy controls, emoji picker, image lightbox, voice message playback, file attachment rendering, link preview cards, theme switching, offline queue indicators, and toast notifications.')

H('Backend Unit & Integration Tests (305)', level=3)
P('Every API endpoint, authentication flow, database operation, Socket.IO event handler, email service, SMS service, and AI integration is tested. Backend tests use an in-memory database (where possible) or a test database with transaction rollback to ensure isolation. The backend test suite covers 73% of the backend codebase.')
P('Key areas covered: user registration and login (all 4 authentication methods), JWT issuance and refresh, friend request lifecycle, group CRUD and role management, message send/receive/edit/delete, reaction and reply creation, file upload validation, rate limiting enforcement, AI response generation, email and SMS delivery, and WebSocket event routing.')

H('End-to-End Browser Tests (156)', level=3)
P('Playwright drives real browsers (Desktop Chrome and iPhone 14 emulation) through complete user journeys. Two authenticated test users interact simultaneously to verify real-time message delivery, typing indicators, and presence updates. E2E tests cover login, registration, conversation list, opening chats, sending and receiving messages, creating groups, sending friend requests, visiting profiles, toggling settings, and verifying dashboard analytics.')
tech('E2E test results are cached to disk (.test-cache/) and displayed in the developer dashboard. A custom Playwright reporter formats results and persists them automatically after each run.')


# ═════════════════════════════════════════════════════════════════════════════
# 13. ANALYTICS DASHBOARD
# ═════════════════════════════════════════════════════════════════════════════
H('Analytics Dashboard')

P('Chatr includes a custom-built project intelligence dashboard that provides real-time visibility into code health, development velocity, test results, security posture, and architecture inventory. The dashboard auto-refreshes every 30 seconds and is accessible from the main navigation.')

img('10-dashboard-top', 'Dashboard \u2014 metric cards, code health gauges, and commit intelligence', width=WIDE_W)

H('Metric Cards', level=2)
P('The top section displays 17+ live metric cards:')
B('Total commits, total lines of code, total source files, total tests.')
B('REST API endpoint count, WebSocket event type count, UI component count, database model count.')
B('Dependency count (frontend + backend), branch count, largest file size.')
P('Each metric is computed in real time by parsing the actual codebase (not cached snapshots), so the numbers are always current.')

H('Code Health Gauges', level=2)
P('Circular gauge charts display code health metrics at a glance: average file size, test coverage by area (frontend, backend, e2e), development velocity (commits per day), and largest-file size. Green indicates healthy, amber indicates caution, and red indicates attention needed.')

H('Commit Intelligence', level=2)
P('The dashboard analyses the Git history and presents:')
B('Change type distribution (features, fixes, refactors, tests, docs) as a breakdown chart.')
B('Commit size distribution (small, medium, large, extra-large) for understanding development patterns.')
B('Weekly commit trend showing velocity over time.')
B('Top 10 largest commits with file counts and descriptions.')
B('Contribution heatmap showing activity by hour and day of week.')

H('Security Audit', level=2)
P('A dedicated section runs npm audit and displays the count and severity of dependency vulnerabilities. Build health status is shown for both frontend (Next.js build) and backend (TypeScript compilation), with pass/fail indicators.')

H('Live Test Runner', level=2)
P('The dashboard includes an embedded test runner that executes frontend, backend, or e2e test suites directly from the browser. Results stream in real time, with filtering by suite, status (pass/fail/skip), and search. Failed tests are highlighted with error messages and stack traces. The runner supports re-running individual failed tests.')

img('09-dashboard-full', 'Full analytics dashboard (scrolled)', width=WIDE_W)


# ═════════════════════════════════════════════════════════════════════════════
# 14. DEVELOPER EXPERIENCE
# ═════════════════════════════════════════════════════════════════════════════
H('Developer Experience')

P('Chatr is not just a product \u2014 it is a developer platform. The tooling built alongside the application demonstrates the same attention to quality as the user-facing features.')

H('Built-in Documentation', level=2)
P('A searchable documentation site is available at /docs within the application. It includes architecture diagrams, setup instructions, API reference, component catalogue, and code examples. The docs are written in a developer-friendly format and kept in sync with the codebase.')

H('Email Template Preview', level=2)
P('All transactional email templates (welcome, verification, password reset, login code, friend request, group invite) are viewable in a dedicated preview page at /email-preview. Developers can see exactly what each email looks like before sending, with sample data pre-filled.')

pair('12-docs', '13-email-templates',
     'Built-in documentation',
     'Email template previews')

H('API Documentation', level=2)
P('A Swagger UI interface provides interactive documentation for all 70+ REST API endpoints. Developers can read descriptions, view request/response schemas, and test endpoints directly from the browser with their authentication token.')

H('Component Demos', level=2)
P('Interactive demonstrations of every UI component are available for visual testing and design reference. This serves as a living style guide that stays in sync with the actual component implementations.')

H('System Log Viewer', level=2)
P('An in-app log viewer with filtering, search, and severity levels (info, warn, error) helps developers diagnose issues in real time without SSH access to the server.')

H('Docker Compose', level=2)
P('A single docker-compose up command provisions PostgreSQL, Redis, the backend server, and the frontend server \u2014 enabling a developer to go from git clone to running application in under 2 minutes.')


# ═════════════════════════════════════════════════════════════════════════════
# 15. BY THE NUMBERS
# ═════════════════════════════════════════════════════════════════════════════
H('By the Numbers')

stat_line('User-facing features', '50+')
stat_line('Automated tests', '1,300+')
stat_line('Lines of code', '78,000+')
stat_line('Source files', '369')
stat_line('REST API endpoints', '70+')
stat_line('WebSocket event types', '40+')
stat_line('UI components', '176 (60+ custom)')
stat_line('Database models', '9')
stat_line('Authentication methods', '4 (email, SMS, 2FA, login code)')
stat_line('Real-time indicators', 'Typing, recording, presence, read receipts, ghost typing')
stat_line('Message types', '7 (text, voice, image, video, file, code block, link preview)')
stat_line('File upload limit', '50 MB')
stat_line('Media support', 'Images, video, audio, PDF, documents, archives')
stat_line('Frontend test coverage', '99%')
stat_line('Backend test coverage', '73%')
stat_line('E2E browsers', 'Desktop Chrome + iPhone 14')
stat_line('Deployment', 'AWS (EC2, RDS, ElastiCache, S3)')
stat_line('Offline support', 'IndexedDB cache + outbound message queue + audio cache')
stat_line('Development time', '22 days')
stat_line('Total commits', '219')


# ═════════════════════════════════════════════════════════════════════════════
# 16. WHY INVEST IN CHATR
# ═════════════════════════════════════════════════════════════════════════════
H('Why Invest in Chatr')

P('This final section consolidates the case for Chatr \u2014 whether you are evaluating it as a product, a technology investment, or a demonstration of engineering capability.')

H('It Is a Complete Product', level=2, numbered=False)
P('Chatr is not a mockup, a tutorial project, or a proof of concept. It is a production-deployed messaging platform with 50+ user-facing features \u2014 real-time messaging, voice notes, video sharing, file attachments, code blocks, link previews, typing indicators, ghost typing, read receipts, emoji reactions, message replies, message editing, message unsending, group chats with role management, friend requests, blocking, user search, an AI chatbot, automatic conversation summaries, toast notifications, offline support, an outbound message queue, audio caching, storage management, dark and light themes, a profile system with avatar and cover image uploads, granular privacy controls, two-factor authentication, SMS verification, email verification, password strength indicators, and a fully embeddable customer support widget with white-label customisation. Every one of these features is built, integrated, and tested. It works today.')

H('It Generates Revenue', level=2, numbered=False)
P('The embeddable chat widget is a direct competitor to Intercom, Drift, and Zendesk Chat. It provides real-time customer support with zero visitor friction, full white-label customisation, and zero recurring SaaS fees. A business with a 10-person support team saves $4,700\u2013$11,900/year compared to Intercom. That makes Chatr a product with real, measurable commercial value \u2014 not just a technical showcase.')

H('It Is Tested Like Enterprise Software', level=2, numbered=False)
P('1,300+ automated tests across three tiers (unit, integration, end-to-end) ensure that every feature works, every edge case is handled, and every deployment is safe. End-to-end tests drive real browsers with two simultaneous test users interacting in real time. 99% frontend coverage and 73% backend coverage exceed the industry average for commercial messaging products. A custom developer dashboard visualises test results, code health, and security posture in real time.')

H('It Is Built on Proven Technology', level=2, numbered=False)
P('React 19, Next.js 16, Node.js, Express, PostgreSQL, Redis, AWS \u2014 the same stack trusted by Slack, Shopify, Netflix, and Uber. Any JavaScript/TypeScript developer can join the project and be productive on day one. There is no proprietary framework, no exotic language, no vendor-specific tooling. Every technology choice maximises the size of the available talent pool.')

H('It Scales Without Rewriting', level=2, numbered=False)
P('PM2 cluster mode utilises all CPU cores. The Socket.IO Redis adapter enables horizontal scaling across multiple server instances. PostgreSQL handles millions of messages. S3 handles unlimited media. Nginx load balances traffic. Scaling Chatr means adding servers, not rewriting architecture.')

H('It Demonstrates Exceptional Engineering Range', level=2, numbered=False)
P('Frontend development. Backend development. Real-time WebSocket infrastructure. AI integration. Cloud deployment. Database design. Security and authentication. Accessibility. Automated testing at three tiers. Developer tooling. A custom analytics dashboard. Email and SMS services. An embeddable third-party widget. All of this was designed, built, tested, documented, and deployed by a single developer in 22 days. That is not just a product \u2014 it is a statement of capability.', bold=True, size=10.5)


# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
n = len([f for f in os.listdir(SS) if f.endswith('.png')]) if os.path.isdir(SS) else 0
print(f'Created: {os.path.abspath(OUT)}')
print(f'  {n} screenshots available in {SS}')
